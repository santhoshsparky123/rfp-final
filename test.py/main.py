from fastapi import FastAPI, UploadFile, File, Form, HTTPException, BackgroundTasks, Query
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Dict, Any, Optional
import json
import os
import uuid
import shutil
from datetime import datetime
from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_openai import ChatOpenAI
from langchain.chains import RetrievalQA
from langchain.agents import Tool
from langchain.prompts import ChatPromptTemplate
from langgraph.graph import StateGraph, END
from langgraph.prebuilt import ToolNode
import wikipedia
import docx
from docx import Document
from docx.shared import Inches, Pt
from docx2pdf import convert
import asyncio
import httpx

# Initialize FastAPI app
app = FastAPI(title="RFP Response Agent API")

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Create necessary directories
os.makedirs("uploads", exist_ok=True)
os.makedirs("company_docs", exist_ok=True)
os.makedirs("outputs", exist_ok=True)
os.makedirs("vector_stores", exist_ok=True)

# Pydantic models
class RFPData(BaseModel):
    success: bool
    rfp_data: Dict[str, Any]

class ToolOutput(BaseModel):
    content: str
    source: str

class AgentResponse(BaseModel):
    rfp_id: str
    section_responses: Dict[str, str]
    full_proposal: str
    download_urls: Dict[str, str]

# Store for processed RFPs
RFP_STORE = {}

# OpenAI API configuration - should be set via environment variables
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "your-api-key")

# Initialize OpenAI client
llm = ChatOpenAI(
    temperature=0.2,
    model="gpt-4-turbo-preview",
    api_key=OPENAI_API_KEY
)

# Document processing functions
def process_document(file_path):
    """Extract text and metadata from uploaded documents (PDF or DOCX)"""
    if file_path.endswith('.pdf'):
        loader = PyPDFLoader(file_path)
        documents = loader.load()
    elif file_path.endswith('.docx'):
        loader = Docx2txtLoader(file_path)
        documents = loader.load()
    else:
        raise ValueError("Unsupported file format. Only PDF and DOCX are supported.")
    
    # Split documents into chunks
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200
    )
    chunks = text_splitter.split_documents(documents)
    
    return chunks

def extract_rfp_structure(file_path):
    """Extract RFP structure and generate structured JSON data"""
    chunks = process_document(file_path)
    combined_text = " ".join([chunk.page_content for chunk in chunks])
    
    # Using LLM to extract structured data from RFP
    prompt = ChatPromptTemplate.from_template(
        """You are an expert at analyzing RFP documents. Extract the structure and key information from the following 
        RFP text and format it as a structured JSON with metadata, sections, questions, and requirements.
        
        The structure should be similar to:
        {
            "success": true,
            "rfp_data": {
                "metadata": {
                    "title": "...",
                    "organization": "...",
                    "issue_date": "...",
                    "due_date": "...",
                    "contact_information": "...",
                    "submission_requirements": "..."
                },
                "sections": [...],
                "questions": [...],
                "requirements": [...]
            }
        }
        
        RFP Text:
        {text}
        
        Respond ONLY with the JSON data, nothing else.
        """
    )
    
    chain = prompt | llm
    response = chain.invoke({"text": combined_text})
    
    # Extract JSON from response
    try:
        # Find JSON in the response text - it might be wrapped in backticks or other formatting
        json_start = response.content.find('{')
        json_end = response.content.rfind('}') + 1
        json_str = response.content[json_start:json_end]
        structured_data = json.loads(json_str)
        return structured_data
    except Exception as e:
        print(f"Error extracting JSON: {e}")
        print(f"Response content: {response.content}")
        raise HTTPException(status_code=500, detail=f"Failed to extract RFP structure: {str(e)}")

# Tool implementations
class CompanyKnowledgeTool:
    """RAG tool for retrieving information from company documents"""
    
    def __init__(self, vector_store_path=None):
        self.embeddings = OpenAIEmbeddings(api_key=OPENAI_API_KEY)
        self.vector_store = None
        if vector_store_path and os.path.exists(vector_store_path):
            self.vector_store = FAISS.load_local(vector_store_path, self.embeddings)
    
    def ingest_documents(self, file_paths):
        """Process and index company documents"""
        all_chunks = []
        for file_path in file_paths:
            chunks = process_document(file_path)
            all_chunks.extend(chunks)
        
        # Create vector store
        if not all_chunks:
            return False
        
        store_id = f"vector_stores/company_docs_{uuid.uuid4()}"
        self.vector_store = FAISS.from_documents(all_chunks, self.embeddings)
        self.vector_store.save_local(store_id)
        return store_id
    
    def query(self, question):
        """Query company knowledge base"""
        if not self.vector_store:
            return {"content": "No company documents have been uploaded yet.", "source": "system"}
        
        retriever = self.vector_store.as_retriever(
            search_type="similarity",
            search_kwargs={"k": 3}
        )
        
        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",
            retriever=retriever,
            return_source_documents=True
        )
        
        result = qa_chain.invoke({"query": question})
        sources = [doc.metadata.get("source", "Unknown") for doc in result["source_documents"]]
        
        return {
            "content": result["result"],
            "source": f"Company documents: {', '.join(set(sources))}"
        }

class WikipediaTool:
    """Tool for retrieving information from Wikipedia"""
    
    def query(self, query):
        """Query Wikipedia API"""
        try:
            # Search for relevant Wikipedia pages
            search_results = wikipedia.search(query, results=3)
            if not search_results:
                return {"content": f"No Wikipedia results found for '{query}'", "source": "Wikipedia"}
            
            # Get the first page (most relevant)
            page_title = search_results[0]
            page = wikipedia.page(page_title)
            
            # Extract a summary
            summary = wikipedia.summary(page_title, sentences=5)
            
            return {
                "content": summary,
                "source": f"Wikipedia: {page.url}"
            }
        except Exception as e:
            return {
                "content": f"Error retrieving information from Wikipedia: {str(e)}",
                "source": "Wikipedia"
            }

class PricingTool:
    """Tool for determining pricing for RFP responses"""
    
    def __init__(self, pricing_data_path=None):
        self.embeddings = OpenAIEmbeddings(api_key=OPENAI_API_KEY)
        self.vector_store = None
        
        # Use default pricing data or provided data
        if pricing_data_path and os.path.exists(pricing_data_path):
            self.load_pricing_data(pricing_data_path)
        else:
            # Default pricing data if none provided
            self.default_pricing = {
                "hourly_rates": {
                    "frontend_developer": 100,
                    "backend_developer": 120,
                    "designer": 90,
                    "project_manager": 110,
                    "qa_engineer": 85
                },
                "project_types": {
                    "e-commerce": {"base_price": 8000, "complexity_multipliers": {"low": 1.0, "medium": 1.5, "high": 2.0}},
                    "content_management": {"base_price": 5000, "complexity_multipliers": {"low": 1.0, "medium": 1.3, "high": 1.8}},
                    "custom_app": {"base_price": 10000, "complexity_multipliers": {"low": 1.0, "medium": 1.5, "high": 2.2}}
                }
            }
    
    def load_pricing_data(self, file_path):
        """Load pricing data from a file"""
        try:
            with open(file_path, 'r') as f:
                self.default_pricing = json.load(f)
        except Exception as e:
            print(f"Error loading pricing data: {e}")
            # Fall back to default pricing
            self.default_pricing = {
                "hourly_rates": {
                    "frontend_developer": 100,
                    "backend_developer": 120,
                    "designer": 90,
                    "project_manager": 110,
                    "qa_engineer": 85
                },
                "project_types": {
                    "e-commerce": {"base_price": 8000, "complexity_multipliers": {"low": 1.0, "medium": 1.5, "high": 2.0}},
                    "content_management": {"base_price": 5000, "complexity_multipliers": {"low": 1.0, "medium": 1.3, "high": 1.8}},
                    "custom_app": {"base_price": 10000, "complexity_multipliers": {"low": 1.0, "medium": 1.5, "high": 2.2}}
                }
            }
    
    def calculate_price(self, rfp_data):
        """Calculate pricing based on RFP requirements"""
        prompt = ChatPromptTemplate.from_template(
            """You are a pricing expert at a software development company. Analyze the following RFP data and 
            calculate an appropriate price quote based on our pricing structure.
            
            RFP Data:
            {rfp_data}
            
            Our Pricing Structure:
            {pricing_structure}
            
            Consider the project scope, timeline, complexity, and requirements when determining the price.
            Provide a detailed breakdown of the costs and the total estimated price.
            """
        )
        
        chain = prompt | llm
        response = chain.invoke({
            "rfp_data": json.dumps(rfp_data, indent=2),
            "pricing_structure": json.dumps(self.default_pricing, indent=2)
        })
        
        return {
            "content": response.content,
            "source": "Pricing Analysis Tool"
        }

# LangGraph Implementation
def create_rfp_agent_graph(rfp_data, company_knowledge_tool, wiki_tool, pricing_tool):
    """Create a LangGraph workflow for RFP response generation"""
    
    # Define the agent state
    class AgentState(BaseModel):
        rfp_data: Dict
        current_question_id: Optional[str] = None
        responses: Dict[str, str] = {}
        intermediate_steps: List = []
        
    # Router node to decide which question to answer next
    def question_router(state):
        """Route to the appropriate question handler or to final compilation"""
        # Check if all questions have been answered
        all_questions = [q["id"] for q in state.rfp_data["rfp_data"]["questions"]]
        answered_questions = list(state.responses.keys())
        
        # If all questions answered, proceed to final compilation
        if set(all_questions).issubset(set(answered_questions)):
            return "compile_final_response"
        
        # Otherwise, pick the next unanswered question
        for q_id in all_questions:
            if q_id not in answered_questions:
                state.current_question_id = q_id
                return "answer_question"
        
        # Fallback
        return "compile_final_response"
    
    # Answer question node
    def answer_question(state):
        """Generate answer for the current question"""
        current_q_id = state.current_question_id
        questions = state.rfp_data["rfp_data"]["questions"]
        
        # Find the current question
        current_question = None
        for q in questions:
            if q["id"] == current_q_id:
                current_question = q
                break
        
        if not current_question:
            return state
        
        # Find related requirements
        related_reqs = []
        if "related_requirements" in current_question:
            req_ids = current_question["related_requirements"]
            for req in state.rfp_data["rfp_data"]["requirements"]:
                if req["id"] in req_ids:
                    related_reqs.append(req)
        
        # Generate answer prompt
        prompt = ChatPromptTemplate.from_template(
            """You are an expert proposal writer responding to an RFP. 
            Answer the following question professionally and persuasively based on our company's knowledge and capabilities.
            
            Question: {question}
            Question Type: {question_type}
            Related Requirements: {related_requirements}
            
            Use the tools available to gather specific information needed to craft a convincing response.
            Format your response according to the requested format: {format}
            {word_limit_instruction}
            
            Your answer should be comprehensive, address all requirements, highlight our company's strengths,
            and demonstrate why we are the best fit for this project.
            """
        )
        
        # Add word limit instruction if specified
        word_limit_instruction = ""
        if current_question.get("word_limit"):
            word_limit_instruction = f"Keep your response within {current_question['word_limit']} words."
        
        # Create the main agent node to answer questions
        tools = [
            Tool(
                name="company_knowledge",
                func=company_knowledge_tool.query,
                description="Use this tool to retrieve information from company documents about our capabilities, experience, and previous projects."
            ),
            Tool(
                name="wikipedia",
                func=wiki_tool.query,
                description="Use this tool to gather general information about technical concepts, standards, or industry knowledge."
            ),
            Tool(
                name="pricing_analysis",
                func=lambda q: pricing_tool.calculate_price(state.rfp_data),
                description="Use this tool to calculate and justify pricing for the proposal based on project requirements."
            )
        ]
        
        # Create the tool-using agent
        tool_node = ToolNode(tools=tools, llm=llm)
        
        # Run the prompt through the tool node
        question_input = {
            "question": current_question["question"],
            "question_type": current_question.get("type", "General"),
            "related_requirements": json.dumps(related_reqs, indent=2) if related_reqs else "None specified",
            "format": current_question.get("format", "Paragraph"),
            "word_limit_instruction": word_limit_instruction
        }
        
        # Generate the agent prompt
        agent_prompt = prompt.format_messages(**question_input)
        
        # Convert prompt messages to a single string for the tool node
        prompt_str = "\n".join([msg.content for msg in agent_prompt])
        
        # Execute the tool node
        result = tool_node.invoke({"messages": [{"role": "user", "content": prompt_str}]})
        
        # Extract the final answer
        answer = result["messages"][-1]["content"]
        
        # Store the answer
        state.responses[current_q_id] = answer
        
        return state
    
    # Compile final response
    def compile_final_response(state):
        """Compile all answers into a final proposal"""
        prompt = ChatPromptTemplate.from_template(
            """You are an expert proposal writer. Compile the following question responses into a cohesive, professional
            proposal document that addresses the original RFP requirements.
            
            RFP Title: {title}
            RFP Organization: {organization}
            
            Question Responses:
            {responses}
            
            Format the proposal with appropriate sections, an executive summary, introduction, and conclusion.
            Ensure the document flows well and presents a compelling case for why our company should be selected.
            
            The final proposal should be in Markdown format with appropriate headings, bullet points, and formatting.
            """
        )
        
        # Format responses for the prompt
        formatted_responses = ""
        for q_id, answer in state.responses.items():
            question_text = "Unknown Question"
            for q in state.rfp_data["rfp_data"]["questions"]:
                if q["id"] == q_id:
                    question_text = q["question"]
                    break
            
            formatted_responses += f"## {question_text}\n\n{answer}\n\n"
        
        # Get metadata
        metadata = state.rfp_data["rfp_data"]["metadata"]
        
        # Generate the final proposal
        chain = prompt | llm
        response = chain.invoke({
            "title": metadata.get("title", "RFP Response"),
            "organization": metadata.get("organization", "Client Organization"),
            "responses": formatted_responses
        })
        
        # Store the final proposal in the state
        state.responses["final_proposal"] = response.content
        
        return state
    
    # Create the graph
    workflow = StateGraph(AgentState)
    
    # Add nodes
    workflow.add_node("router", question_router)
    workflow.add_node("answer_question", answer_question)
    workflow.add_node("compile_final_response", compile_final_response)
    
    # Add edges
    workflow.add_edge("router", "answer_question")
    workflow.add_edge("answer_question", "router")
    workflow.add_edge("router", "compile_final_response")
    workflow.add_edge("compile_final_response", END)
    
    # Set the entry point
    workflow.set_entry_point("router")
    
    # Compile the graph
    return workflow.compile()

def generate_proposal_document(rfp_id, responses):
    """Generate a Word document and PDF from proposal responses"""
    # Create Word document
    doc = Document()
    
    # Add title
    title = responses.get("title", "RFP Response")
    doc.add_heading(title, 0)
    
    # Add date
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph("\n")
    
    # Add content
    final_proposal = responses.get("final_proposal", "")
    if final_proposal:
        # Split the markdown content by headings
        sections = final_proposal.split("## ")
        
        # Add each section
        for i, section in enumerate(sections):
            if i == 0 and not section.startswith("#"):  # Introduction part
                doc.add_paragraph(section)
                continue
                
            # Split heading from content
            lines = section.split("\n", 1)
            if len(lines) > 0:
                heading = lines[0]
                doc.add_heading(heading, level=2)
                
                if len(lines) > 1:
                    content = lines[1]
                    paragraphs = content.split("\n\n")
                    for para in paragraphs:
                        if para.strip():
                            # Check if it's a bullet point list
                            if para.strip().startswith("- "):
                                bullet_items = para.split("- ")
                                for item in bullet_items:
                                    if item.strip():
                                        doc.add_paragraph(item.strip(), style='List Bullet')
                            else:
                                doc.add_paragraph(para)
    else:
        # Add individual question responses
        for q_id, response in responses.items():
            if q_id != "final_proposal" and q_id != "title":
                doc.add_heading(f"Question {q_id}", level=2)
                doc.add_paragraph(response)
    
    # Create file paths
    docx_path = f"outputs/{rfp_id}_proposal.docx"
    pdf_path = f"outputs/{rfp_id}_proposal.pdf"
    
    # Save Word document
    doc.save(docx_path)
    
    # Convert to PDF
    try:
        convert(docx_path, pdf_path)
    except Exception as e:
        print(f"Error converting to PDF: {e}")
        # PDF conversion failed, but we still have the DOCX
    
    # Return paths
    return {
        "docx": docx_path,
        "pdf": pdf_path if os.path.exists(pdf_path) else None
    }

# API Endpoints
@app.post("/upload-rfp/", response_model=dict)
async def upload_rfp(file: UploadFile = File(...)):
    """Upload and process an RFP document"""
    # Generate unique ID for this RFP
    rfp_id = str(uuid.uuid4())
    
    # Save the uploaded file
    file_extension = os.path.splitext(file.filename)[1]
    file_path = f"uploads/rfp_{rfp_id}{file_extension}"
    
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    try:
        # Extract structured data from RFP
        structured_data = extract_rfp_structure(file_path)
        
        # Store the processed RFP
        RFP_STORE[rfp_id] = {
            "file_path": file_path,
            "structured_data": structured_data,
            "company_docs": [],
            "responses": {}
        }
        
        return {
            "rfp_id": rfp_id,
            "message": "RFP uploaded and processed successfully",
            "structured_data": structured_data
        }
    except Exception as e:
        # Clean up on error
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(status_code=500, detail=f"Error processing RFP: {str(e)}")

@app.post("/upload-company-docs/{rfp_id}", response_model=dict)
async def upload_company_docs(rfp_id: str, files: List[UploadFile] = File(...)):
    """Upload company documents for RAG"""
    if rfp_id not in RFP_STORE:
        raise HTTPException(status_code=404, detail="RFP not found")
    
    saved_files = []
    try:
        for file in files:
            # Save the uploaded file
            file_extension = os.path.splitext(file.filename)[1]
            file_path = f"company_docs/{rfp_id}_{file.filename}"
            
            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
            
            saved_files.append(file_path)
        
        # Process and index the documents
        company_tool = CompanyKnowledgeTool()
        vector_store_id = company_tool.ingest_documents(saved_files)
        
        # Update RFP store
        RFP_STORE[rfp_id]["company_docs"].extend(saved_files)
        RFP_STORE[rfp_id]["vector_store_id"] = vector_store_id
        
        return {
            "rfp_id": rfp_id,
            "message": f"{len(saved_files)} company documents uploaded and processed",
            "vector_store_id": vector_store_id
        }
    except Exception as e:
        # Clean up on error
        for file_path in saved_files:
            if os.path.exists(file_path):
                os.remove(file_path)
        raise HTTPException(status_code=500, detail=f"Error processing company documents: {str(e)}")

@app.post("/generate-response/{rfp_id}", response_model=dict)
async def generate_response(
    rfp_id: str, 
    background_tasks: BackgroundTasks,
    custom_pricing_path: Optional[str] = Form(None)
):
    """Generate RFP response using the agent graph"""
    if rfp_id not in RFP_STORE:
        raise HTTPException(status_code=404, detail="RFP not found")
    
    rfp_data = RFP_STORE[rfp_id]
    
    # Initialize tools
    company_tool = CompanyKnowledgeTool(rfp_data.get("vector_store_id"))
    wiki_tool = WikipediaTool()
    pricing_tool = PricingTool(custom_pricing_path)
    
    # Create task ID for tracking
    task_id = str(uuid.uuid4())
    
    # Function to run in background
    async def process_rfp_response():
        try:
            # Create agent graph
            agent_graph = create_rfp_agent_graph(
                rfp_data["structured_data"],
                company_tool,
                wiki_tool,
                pricing_tool
            )
            
            # Initialize state
            initial_state = {
                "rfp_data": rfp_data["structured_data"],
                "responses": {},
                "intermediate_steps": []
            }
            
            # Execute the graph
            result = agent_graph.invoke(initial_state)
            
            # Get the final responses
            responses = result["responses"]
            
            # Add title from RFP
            if "metadata" in rfp_data["structured_data"]["rfp_data"]:
                responses["title"] = rfp_data["structured_data"]["rfp_data"]["metadata"].get("title", "RFP Response")
            
            # Generate documents
            doc_paths = generate_proposal_document(rfp_id, responses)
            
            # Update RFP store with responses
            RFP_STORE[rfp_id]["responses"] = responses
            RFP_STORE[rfp_id]["document_paths"] = doc_paths
            RFP_STORE[rfp_id]["status"] = "completed"
            
        except Exception as e:
            print(f"Error generating response: {e}")
            RFP_STORE[rfp_id]["status"] = "failed"
            RFP_STORE[rfp_id]["error"] = str(e)
    
    # Start background processing
    background_tasks.add_task(process_rfp_response)
    
    # Update status to processing
    RFP_STORE[rfp_id]["status"] = "processing"
    
    return {
        "rfp_id": rfp_id,
        "message": "RFP response generation started",
        "status": "processing"
    }

@app.get("/response-status/{rfp_id}", response_model=dict)
async def check_response_status(rfp_id: str):
    """Check the status of an RFP response generation"""
    if rfp_id not in RFP_STORE:
        raise HTTPException(status_code=404, detail="RFP not found")
    
    rfp_data = RFP_STORE[rfp_id]
    status = rfp_data.get("status", "not_started")
    
    response = {
        "rfp_id": rfp_id,
        "status": status
    }
    
    if status == "completed":
        # Add URLs for downloading documents
        doc_paths = rfp_data.get("document_paths", {})
        base_url = "/download-document"
        
        download_urls = {}
        for doc_type, path in doc_paths.items():
            if path:
                download_urls[doc_type] = f"{base_url}/{rfp_id}/{doc_type}"
        
        response["download_urls"] = download_urls
    elif status == "failed":
        response["error"] = rfp_data.get("error", "Unknown error")
    
    return response

@app.get("/download-document/{rfp_id}/{doc_type}")
async def download_document(rfp_id: str, doc_type: str):
    """Download generated documents"""
    if rfp_id not in RFP_STORE:
        raise HTTPException(status_code=404, detail="RFP not found")
    
    rfp_data = RFP_STORE[rfp_id]
    if "document_paths" not in rfp_data:
        raise HTTPException(status_code=404, detail="Documents not found")
    
    doc_paths = rfp_data["document_paths"]
    if doc_type not in doc_paths or not doc_paths[doc_type]:
        raise HTTPException(status_code=404, detail=f"{doc_type} document not available")
    
    file_path = doc_paths[doc_type]
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Document file not found")
    
    # Set the appropriate content type
    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if doc_type == "docx" else "application/pdf"
    
    return FileResponse(path=file_path, media_type=media_type, filename=os.path.basename(file_path))

@app.get("/rfp-questions/{rfp_id}", response_model=List[dict])
async def get_rfp_questions(rfp_id: str):
    """Get questions from the processed RFP"""
    if rfp_id not in RFP_STORE:
        raise HTTPException(status_code=404, detail="RFP not found")
    
    rfp_data = RFP_STORE[rfp_id]
    if "structured_data" not in rfp_data or "rfp_data" not in rfp_data["structured_data"]:
        raise HTTPException(status_code=404, detail="RFP questions not found")
    
    questions = rfp_data["structured_data"]["rfp_data"].get("questions", [])
    return questions

@app.get("/all-rfps/", response_model=List[dict])
async def list_all_rfps():
    """List all processed RFPs"""
    rfps = []
    for rfp_id, data in RFP_STORE.items():
        rfp_info = {
            "rfp_id": rfp_id,
            "title": data.get("structured_data", {}).get("rfp_data", {}).get("metadata", {}).get("title", "Untitled RFP"),
            "status": data.get("status", "not_started"),
            "upload_date": datetime.fromtimestamp(os.path.getctime(data["file_path"])).strftime('%Y-%m-%d %H:%M:%S')
        }
        rfps.append(rfp_info)
    
@app.post("/edit-response/{rfp_id}/{question_id}")
async def edit_response(rfp_id: str, question_id: str, response: str = Form(...)):
    """Manually edit a response to an RFP question"""
    if rfp_id not in RFP_STORE:
        raise HTTPException(status_code=404, detail="RFP not found")
    
    rfp_data = RFP_STORE[rfp_id]
    if "responses" not in rfp_data:
        rfp_data["responses"] = {}
    
    # Update the specific response
    rfp_data["responses"][question_id] = response
    
    # If final_proposal exists, mark it as outdated
    if "final_proposal" in rfp_data["responses"]:
        rfp_data["final_proposal_outdated"] = True
    
    return {
        "rfp_id": rfp_id,
        "question_id": question_id,
        "message": "Response updated successfully"
    }

@app.post("/regenerate-final-proposal/{rfp_id}")
async def regenerate_final_proposal(rfp_id: str):
    """Regenerate the final proposal after editing individual responses"""
    if rfp_id not in RFP_STORE:
        raise HTTPException(status_code=404, detail="RFP not found")
    
    rfp_data = RFP_STORE[rfp_id]
    if "responses" not in rfp_data or not rfp_data["responses"]:
        raise HTTPException(status_code=400, detail="No responses found to generate final proposal")
    
    # Create prompt for final proposal generation
    prompt = ChatPromptTemplate.from_template(
        """You are an expert proposal writer. Compile the following question responses into a cohesive, professional
        proposal document that addresses the original RFP requirements.
        
        RFP Title: {title}
        RFP Organization: {organization}
        
        Question Responses:
        {responses}
        
        Format the proposal with appropriate sections, an executive summary, introduction, and conclusion.
        Ensure the document flows well and presents a compelling case for why our company should be selected.
        
        The final proposal should be in Markdown format with appropriate headings, bullet points, and formatting.
        """
    )
    
    # Format responses for the prompt
    formatted_responses = ""
    for q_id, answer in rfp_data["responses"].items():
        if q_id != "final_proposal" and q_id != "title":
            question_text = "Unknown Question"
            if "structured_data" in rfp_data and "rfp_data" in rfp_data["structured_data"]:
                for q in rfp_data["structured_data"]["rfp_data"].get("questions", []):
                    if q["id"] == q_id:
                        question_text = q["question"]
                        break
            
            formatted_responses += f"## {question_text}\n\n{answer}\n\n"
    
    # Get metadata
    metadata = rfp_data.get("structured_data", {}).get("rfp_data", {}).get("metadata", {})
    
    # Generate the final proposal
    chain = prompt | llm
    response = chain.invoke({
        "title": metadata.get("title", "RFP Response"),
        "organization": metadata.get("organization", "Client Organization"),
        "responses": formatted_responses
    })
    
    # Update the final proposal
    rfp_data["responses"]["final_proposal"] = response.content
    rfp_data["final_proposal_outdated"] = False
    
    # Regenerate documents
    if "responses" in rfp_data and "title" in metadata:
        rfp_data["responses"]["title"] = metadata.get("title", "RFP Response")
    doc_paths = generate_proposal_document(rfp_id, rfp_data["responses"])
    rfp_data["document_paths"] = doc_paths
    
    # Add URLs for downloading documents
    base_url = "/download-document"
    download_urls = {}
    for doc_type, path in doc_paths.items():
        if path:
            download_urls[doc_type] = f"{base_url}/{rfp_id}/{doc_type}"
    
@app.get("/rfp-data/{rfp_id}")
async def get_rfp_data(rfp_id: str):
    """Get the full structured RFP data"""
    if rfp_id not in RFP_STORE:
        raise HTTPException(status_code=404, detail="RFP not found")
    
    return RFP_STORE[rfp_id]["structured_data"]

@app.get("/response/{rfp_id}/{question_id}")
async def get_response(rfp_id: str, question_id: str):
    """Get a specific response for an RFP question"""
    if rfp_id not in RFP_STORE:
        raise HTTPException(status_code=404, detail="RFP not found")
    
    rfp_data = RFP_STORE[rfp_id]
    if "responses" not in rfp_data or question_id not in rfp_data["responses"]:
        raise HTTPException(status_code=404, detail="Response not found")
    
    return {
        "rfp_id": rfp_id,
        "question_id": question_id,
        "response": rfp_data["responses"][question_id]
    }

@app.get("/final-proposal/{rfp_id}")
async def get_final_proposal(rfp_id: str):
    """Get the final proposal text"""
    if rfp_id not in RFP_STORE:
        raise HTTPException(status_code=404, detail="RFP not found")
    
    rfp_data = RFP_STORE[rfp_id]
    if "responses" not in rfp_data or "final_proposal" not in rfp_data["responses"]:
        raise HTTPException(status_code=404, detail="Final proposal not found")
    
    return {
        "rfp_id": rfp_id,
        "proposal": rfp_data["responses"]["final_proposal"],
        "is_outdated": rfp_data.get("final_proposal_outdated", False)
    }

# Health check endpoint
@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "ok", "version": "1.0.0"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)